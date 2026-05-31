import io
import logging
import re
from datetime import date
from django.conf import settings
from docx import Document
from core.agent.domain.tools import BaseTool, ToolResult
from core.agent.infrastructure.gemini_adapter import GeminiAdapter
from core.agent.infrastructure.rag_utils import get_rag_context

logger = logging.getLogger(__name__)

DOC_TYPES = {
    'propuesta': 'una propuesta de negocio profesional',
    'contrato': 'un contrato de servicios profesional',
    'informe': 'un informe de resultados ejecutivo',
    'brief': 'un brief creativo o de proyecto',
    'presupuesto': 'un presupuesto detallado de servicios',
}


def _markdown_to_docx(doc: Document, text: str) -> None:
    """Convierte markdown simple (# ## ### **texto**) a párrafos Word."""
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph('')
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        else:
            # Parsear spans **bold** y renderizarlos correctamente en DOCX
            paragraph = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    paragraph.add_run(part[2:-2]).bold = True
                elif part:
                    paragraph.add_run(part)


def _markdown_to_pdf(text: str) -> bytes:
    """Convierte markdown simple (# ## ### **bold**) a PDF con fpdf2."""
    from fpdf import FPDF  # noqa: PLC0415 — lazy import, fpdf2 solo disponible en backend
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(20, 20, 20)
    pdf.add_page()

    for line in text.split('\n'):
        line = line.strip()
        if not line:
            pdf.ln(4)
            continue
        if line.startswith('### '):
            pdf.set_font('Helvetica', 'B', 13)
            pdf.multi_cell(0, 8, line[4:])
            pdf.ln(2)
        elif line.startswith('## '):
            pdf.set_font('Helvetica', 'B', 15)
            pdf.multi_cell(0, 10, line[3:])
            pdf.ln(3)
        elif line.startswith('# '):
            pdf.set_font('Helvetica', 'B', 18)
            pdf.multi_cell(0, 12, line[2:])
            pdf.ln(4)
        else:
            parts = re.split(r'(\*\*.*?\*\*)', line)
            pdf.set_x(pdf.l_margin)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    pdf.set_font('Helvetica', 'B', 11)
                    pdf.write(6, part[2:-2])
                else:
                    pdf.set_font('Helvetica', '', 11)
                    if part:
                        pdf.write(6, part)
            pdf.ln(6)

    return bytes(pdf.output())


class GenerateDocumentTool(BaseTool):
    name = 'generate_document'

    def __init__(self):
        self._gemini = GeminiAdapter()
        self._api_key = settings.GEMINI_API_KEY
        self._model = getattr(settings, 'AI_MODEL', 'gemini-2.5-flash')

    def execute(self, doc_type: str, description: str) -> ToolResult:
        doc_type = doc_type.lower().strip()
        if doc_type not in DOC_TYPES:
            valid = ', '.join(DOC_TYPES.keys())
            return self._error(f'Tipo de documento no válido. Usa uno de: {valid}')
        if not self._api_key:
            return self._error('API key de Gemini no configurada.')

        type_label = DOC_TYPES[doc_type]
        rag_context = get_rag_context(description)
        rag_section = (
            f'\n\nContexto de tus documentos de negocio:\n{rag_context}'
        ) if rag_context else ''

        prompt = (
            f'Eres un asistente de negocios experto. Redacta {type_label} sobre:\n'
            f'"{description}"'
            f'{rag_section}\n\n'
            f'Estructura el documento con secciones claras usando markdown (# para título, '
            f'## para secciones, ### para subsecciones). '
            f'Usa lenguaje profesional en español. '
            f'El documento debe ser completo y listo para entregar a un cliente.'
        )
        try:
            content = self._gemini.generate_response(
                prompt=prompt, api_key=self._api_key, model_name=self._model
            )
        except Exception as e:
            logger.error(f'Error generando contenido con Gemini: {e}', exc_info=True)
            return self._error(f'Error al generar el documento: {e}')

        doc = Document()
        _markdown_to_docx(doc, content)
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        pdf_bytes = _markdown_to_pdf(content)

        filename = f'{doc_type}_{date.today().isoformat()}.docx'
        pdf_filename = f'{doc_type}_{date.today().isoformat()}.pdf'
        return ToolResult(
            content=f'Documentos "{filename}" y "{pdf_filename}" generados correctamente.',
            tool_name=self.name,
            success=True,
            metadata={
                'docx_bytes': docx_bytes,
                'filename': filename,
                'pdf_bytes': pdf_bytes,
                'pdf_filename': pdf_filename,
                'doc_type': doc_type,
            },
        )
