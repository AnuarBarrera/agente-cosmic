import io
import logging
import re
from datetime import date
from bs4 import BeautifulSoup
from playwright.sync_api import sync_playwright
from docx import Document
from django.conf import settings
from core.agent.domain.tools import BaseTool, ToolResult
from core.agent.infrastructure.gemini_adapter import GeminiAdapter, FALLBACK_MESSAGE
from core.agent.infrastructure import scrape_guard

logger = logging.getLogger(__name__)

MAX_CONTENT_CHARS = 3000


def _scrape_url(url: str) -> str:
    try:
        with sync_playwright() as pw:
            browser = pw.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(url, timeout=20000, wait_until='domcontentloaded')
            html = page.content()
            browser.close()
        soup = BeautifulSoup(html, 'html.parser')
        for tag in soup(['script', 'style', 'nav', 'footer']):
            tag.decompose()
        return soup.get_text(separator=' ')[:MAX_CONTENT_CHARS]
    except Exception as e:
        logger.warning(f'Error scrapeando {url}: {e}')
        return ''


def _markdown_to_docx(doc: Document, text: str) -> None:
    for line in text.split('\n'):
        line = line.strip()
        if not line:
            doc.add_paragraph('')
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        else:
            paragraph = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    paragraph.add_run(part[2:-2]).bold = True
                elif part:
                    paragraph.add_run(part)


class ProspectResearchTool(BaseTool):
    name = 'prospect_research'

    def __init__(self):
        self._gemini = GeminiAdapter()

    def execute(self, name: str, urls: list[str]) -> ToolResult:
        if not name.strip():
            return self._error('Debes indicar el nombre del prospecto.')
        api_key = settings.GEMINI_API_KEY
        if not api_key:
            return self._error('API key de Gemini no configurada.')
        model = getattr(settings, 'AI_MODEL', 'gemini-2.5-flash')

        scraped_sections = []
        for url in urls:
            raw = _scrape_url(url)
            if raw:
                safe = scrape_guard.safe_external_content(raw, source=url)
                scraped_sections.append(f'Fuente: {url}\n{safe}')

        web_context = '\n\n'.join(scraped_sections) if scraped_sections else 'No se proporcionaron URLs.'

        prompt = (
            f'Eres un consultor de ventas experto. Analiza la siguiente información sobre el prospecto '
            f'"{name}" y genera un brief de ventas en español.\n\n'
            f'Información recopilada:\n{web_context}\n\n'
            f'El brief debe incluir:\n'
            f'# Brief de Ventas — {name}\n\n'
            f'## 1. Fortalezas\n'
            f'Qué hace bien esta marca/persona.\n\n'
            f'## 2. Áreas de Oportunidad\n'
            f'Dónde puede mejorar.\n\n'
            f'## 3. Ángulo de Venta Recomendado\n'
            f'Cómo presentar nuestra propuesta.\n\n'
            f'## 4. Tono Sugerido\n'
            f'Cómo comunicarse con este prospecto.\n\n'
            f'Sé específico y accionable. Usa datos de la información proporcionada. '
            f'Responde únicamente con el documento en markdown, sin introducción.'
        )
        try:
            brief = self._gemini.generate_response(
                prompt=prompt, api_key=api_key, model_name=model, thinking_budget=0
            )
            if brief == FALLBACK_MESSAGE or not brief.strip():
                return self._error('El servicio de IA no está disponible temporalmente.')

            doc = Document()
            _markdown_to_docx(doc, brief)
            buf = io.BytesIO()
            doc.save(buf)
            docx_bytes = buf.getvalue()

            slug = re.sub(r'[^a-z0-9]+', '_', name.lower())[:30]
            filename = f'brief_{slug}_{date.today().isoformat()}.docx'

            return ToolResult(
                content=f'Brief de *{name}* generado.',
                tool_name=self.name,
                success=True,
                metadata={
                    'docx_bytes': docx_bytes,
                    'filename': filename,
                    'name': name,
                    'urls_scraped': len(scraped_sections),
                },
            )
        except Exception as e:
            logger.error(f'Error generando brief con Gemini: {e}', exc_info=True)
            return self._error(f'Error al generar el brief: {e}')
